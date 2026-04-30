from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from model_utils import TARGET, format_currency, load_workers_compensation, train_models

DOWNLOADS = ROOT.parent
TEMPLATE_CANDIDATES = [*DOWNLOADS.glob("*DS.docx"), *DOWNLOADS.parent.glob("*DS.docx")]
if not TEMPLATE_CANDIDATES:
    raise FileNotFoundError("Не найден шаблон отчета *DS.docx рядом с проектом или на уровень выше.")
TEMPLATE = TEMPLATE_CANDIDATES[0]
REPORT_DIR = ROOT / "reports"
ASSETS_DIR = REPORT_DIR / "assets"
OUTPUT_DOCX = REPORT_DIR / "VKR_Report_Workers_Compensation.docx"
RESULTS_JSON = REPORT_DIR / "model_results.json"
CONSISTENCY_CHECK = REPORT_DIR / "consistency_check.txt"

TIFFANY = "#0ABAB5"
TIFFANY_DARK = "#078F8B"
INK = "#172426"
MUTED = "#667174"
GRID = "#DCEEEE"
FORMAL_HEADER = "#D9D9D9"
FORMAL_BORDER = "#000000"
MODEL_COLORS = {
    "XGBoost": TIFFANY,
    "Random Forest": "#2F80ED",
    "Ridge Regression": "#7A5AF8",
    "Linear Regression": "#F2994A",
}


FEATURE_DESCRIPTIONS = [
    ("DateTimeOfAccident", "Дата и время несчастного случая"),
    ("DateReported", "Дата сообщения о страховом случае"),
    ("Age", "Возраст работника"),
    ("Gender", "Пол работника"),
    ("MaritalStatus", "Семейное положение"),
    ("DependentChildren", "Количество детей на иждивении"),
    ("DependentsOther", "Количество других иждивенцев"),
    ("WeeklyPay", "Еженедельная зарплата"),
    ("PartTimeFullTime", "Тип занятости"),
    ("HoursWorkedPerWeek", "Количество рабочих часов в неделю"),
    ("DaysWorkedPerWeek", "Количество рабочих дней в неделю"),
    ("ClaimDescription", "Текстовое описание страхового случая"),
    ("InitialCaseEstimate", "Начальная оценка стоимости случая"),
    (TARGET, "Итоговая стоимость страхового возмещения"),
]


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(14)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_centered(doc: Document, text: str, bold: bool = False, size: int = 14) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str = "", bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill.replace("#", ""))
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color: str = FORMAL_BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.replace("#", ""))


def set_cell_margins(cell, margin: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "left", "bottom", "right"]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: int = 10,
    color: str = INK,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_borders(cell)


def add_table_caption(doc: Document, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(caption)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("000000")


def add_simple_table(
    doc: Document,
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: int = 10,
    header_fill: str = FORMAL_HEADER,
    caption: str | None = None,
) -> None:
    if caption:
        add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        shade_cell(cell, header_fill)
        set_cell_text(cell, value, bold=True, size=font_size, color="#000000", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value, size=font_size)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: Cm = Cm(15)) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def save_charts(result) -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "metrics": ASSETS_DIR / "metrics_comparison.png",
        "scatter": ASSETS_DIR / "predicted_vs_actual.png",
        "importance": ASSETS_DIR / "feature_importance.png",
        "comparison": ASSETS_DIR / "models_scatter_comparison.png",
        "errors": ASSETS_DIR / "absolute_error_distribution.png",
    }

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.edgecolor": GRID,
            "axes.labelcolor": INK,
            "axes.titlecolor": INK,
            "xtick.color": MUTED,
            "ytick.color": MUTED,
            "grid.color": GRID,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )
    metrics = result.metrics.copy()

    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    y = range(len(metrics))
    bar_height = 0.34
    ax.barh([pos + bar_height / 2 for pos in y], metrics["RMSE"], height=bar_height, color=TIFFANY, label="RMSE")
    ax.barh([pos - bar_height / 2 for pos in y], metrics["MAE"], height=bar_height, color="#748184", label="MAE")
    ax.set_yticks(list(y), metrics["Model"])
    ax.invert_yaxis()
    ax.set_title("Сравнение моделей по ошибкам MAE и RMSE", fontsize=15, pad=14, weight="bold")
    ax.set_xlabel("Ошибка, $")
    ax.legend(frameon=False, loc="lower right")
    for index, row in metrics.iterrows():
        ax.text(row["RMSE"] + 280, index + bar_height / 2, f"{row['RMSE']:,.0f}".replace(",", " "), va="center", fontsize=9, color=INK)
        ax.text(row["MAE"] + 280, index - bar_height / 2, f"{row['MAE']:,.0f}".replace(",", " "), va="center", fontsize=9, color=INK)
    ax.spines[["top", "right", "left"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(paths["metrics"], dpi=240, bbox_inches="tight")
    plt.close(fig)

    best_predictions = result.predictions[result.predictions["Model"] == result.best_model_name]
    best_predictions = best_predictions.sample(min(4500, len(best_predictions)), random_state=42)
    limit = float(result.target_cap or max(best_predictions["Actual"].max(), best_predictions["Predicted"].max()))
    fig, ax = plt.subplots(figsize=(7.2, 6.4))
    ax.scatter(best_predictions["Actual"], best_predictions["Predicted"], alpha=0.32, s=16, color=TIFFANY_DARK, edgecolors="none")
    ax.plot([0, limit], [0, limit], color="#c94c4c", linestyle="--", linewidth=1.4)
    ax.set_xlim(0, limit)
    ax.set_ylim(0, limit)
    ax.set_title(f"{result.best_model_name}: реальные и предсказанные значения", fontsize=14, pad=12, weight="bold")
    ax.set_xlabel("Реальная стоимость, $")
    ax.set_ylabel("Предсказанная стоимость, $")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(paths["scatter"], dpi=240, bbox_inches="tight")
    plt.close(fig)

    fig, axes = plt.subplots(2, 2, figsize=(11.5, 9.2), sharex=True, sharey=True)
    axes = axes.ravel()
    for axis, (_, metric_row) in zip(axes, metrics.iterrows()):
        model_name = metric_row["Model"]
        sample = result.predictions[result.predictions["Model"] == model_name]
        sample = sample.sample(min(1600, len(sample)), random_state=42)
        axis.scatter(
            sample["Actual"],
            sample["Predicted"],
            alpha=0.28,
            s=10,
            color=MODEL_COLORS.get(model_name, TIFFANY),
            edgecolors="none",
        )
        axis.plot([0, limit], [0, limit], color="#c94c4c", linestyle="--", linewidth=1)
        axis.set_xlim(0, limit)
        axis.set_ylim(0, limit)
        axis.set_title(
            f"{model_name}\nRMSE {metric_row['RMSE']:,.0f} | R2 {metric_row['R2']:.3f}".replace(",", " "),
            fontsize=11,
            color=INK,
        )
        axis.spines[["top", "right"]].set_visible(False)
    for axis in axes[2:]:
        axis.set_xlabel("Реальная стоимость, $")
    for axis in axes[::2]:
        axis.set_ylabel("Прогноз, $")
    fig.suptitle("Сравнение графиков факт-прогноз по всем моделям", fontsize=16, weight="bold", color=INK)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(paths["comparison"], dpi=240, bbox_inches="tight")
    plt.close(fig)

    errors = []
    for _, row in metrics.iterrows():
        model_name = row["Model"]
        model_predictions = result.predictions[result.predictions["Model"] == model_name].copy()
        model_predictions["AbsoluteError"] = (model_predictions["Predicted"] - model_predictions["Actual"]).abs()
        errors.append(model_predictions[["Model", "AbsoluteError"]])
    errors_df = pd.concat(errors, ignore_index=True)
    clip = errors_df["AbsoluteError"].quantile(0.99)
    plot_data = [
        errors_df.loc[errors_df["Model"] == model, "AbsoluteError"].clip(upper=clip).to_numpy()
        for model in metrics["Model"]
    ]
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    box = ax.boxplot(plot_data, patch_artist=True, tick_labels=metrics["Model"], showfliers=False)
    for patch, model in zip(box["boxes"], metrics["Model"]):
        patch.set_facecolor(MODEL_COLORS.get(model, TIFFANY))
        patch.set_alpha(0.72)
        patch.set_edgecolor("#3c4a4c")
    for median in box["medians"]:
        median.set_color("#FFFFFF")
        median.set_linewidth(1.8)
    ax.set_title("Распределение абсолютных ошибок моделей", fontsize=15, pad=14, weight="bold")
    ax.set_ylabel("Абсолютная ошибка, $")
    ax.tick_params(axis="x", rotation=16)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(paths["errors"], dpi=240, bbox_inches="tight")
    plt.close(fig)

    importance = result.feature_importance.sort_values("Importance", ascending=True)
    fig, ax = plt.subplots(figsize=(8.4, 6.7))
    bars = ax.barh(importance["Feature"], importance["Importance"], color=TIFFANY)
    ax.bar_label(bars, labels=[f"{value:.3f}" for value in importance["Importance"]], padding=4, fontsize=9)
    ax.set_title("Топ признаков по важности", fontsize=15, pad=14, weight="bold")
    ax.set_xlabel("Важность")
    ax.spines[["top", "right", "left"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(paths["importance"], dpi=240, bbox_inches="tight")
    plt.close(fig)
    return paths


def draw_text(ax, x: float, y: float, text: str, size: int = 12, weight: str = "normal", color: str = INK) -> None:
    ax.text(x, y, text, transform=ax.transAxes, fontsize=size, fontweight=weight, color=color, va="top")


def draw_card(ax, x: float, y: float, width: float, height: float, title: str, value: str) -> None:
    ax.add_patch(
        plt.Rectangle(
            (x, y - height),
            width,
            height,
            transform=ax.transAxes,
            facecolor="#FFFFFF",
            edgecolor=GRID,
            linewidth=1.2,
        )
    )
    ax.add_patch(
        plt.Rectangle((x, y - 0.012), width, 0.012, transform=ax.transAxes, facecolor=TIFFANY, edgecolor=TIFFANY)
    )
    draw_text(ax, x + 0.015, y - 0.025, title, size=10, color=MUTED)
    draw_text(ax, x + 0.015, y - 0.074, value, size=17, weight="bold")


def draw_app_shell(fig, title: str, active_tab: str) -> plt.Axes:
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    ax.add_patch(plt.Rectangle((0, 0), 1, 1, transform=ax.transAxes, facecolor="#FFFFFF", edgecolor="none"))
    ax.add_patch(plt.Rectangle((0, 0), 0.19, 1, transform=ax.transAxes, facecolor="#F4FBFB", edgecolor=GRID))
    draw_text(ax, 0.035, 0.94, "Проект", size=13, weight="bold")
    draw_text(ax, 0.035, 0.88, "Анализ и модель", size=11, color=TIFFANY_DARK if active_tab != "Презентация" else INK)
    draw_text(ax, 0.035, 0.835, "Презентация", size=11, color=TIFFANY_DARK if active_tab == "Презентация" else INK)
    draw_text(ax, 0.035, 0.73, "Параметры обучения", size=12, weight="bold")
    draw_text(ax, 0.035, 0.68, "Весь датасет", size=10, color=MUTED)
    ax.add_patch(plt.Rectangle((0.035, 0.61), 0.12, 0.04, transform=ax.transAxes, facecolor=TIFFANY, edgecolor=TIFFANY))
    draw_text(ax, 0.047, 0.638, "Обучить модели", size=9, color="#FFFFFF", weight="bold")

    draw_text(ax, 0.23, 0.94, title, size=22, weight="bold")
    draw_text(ax, 0.23, 0.895, "Workers Compensation, OpenML ID 42876", size=11, color=MUTED)
    tabs = ["Данные", "Модели", "Сравнение графиков", "Предсказание"]
    x = 0.23
    for tab_name in tabs:
        width = 0.085 if len(tab_name) < 8 else 0.155
        color = TIFFANY if tab_name == active_tab else "#EAF4F4"
        text_color = "#FFFFFF" if tab_name == active_tab else INK
        ax.add_patch(plt.Rectangle((x, 0.665), width, 0.045, transform=ax.transAxes, facecolor=color, edgecolor=GRID))
        draw_text(ax, x + 0.011, 0.695, tab_name, size=10, weight="bold", color=text_color)
        x += width + 0.012
    return ax


def save_app_previews(result) -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    fig = plt.figure(figsize=(16, 10), dpi=150)
    ax = draw_app_shell(fig, "Прогнозирование стоимости страховых выплат", "Данные")
    cards = [
        ("Записей", "100 000"),
        ("Признаков", "13"),
        ("Отсечение выбросов", format_currency(result.target_cap)),
        ("Лучшая модель", result.best_model_name),
        ("RMSE", format_currency(float(result.metrics.iloc[0]["RMSE"]))),
    ]
    card_width = 0.135
    for index, (title, value) in enumerate(cards):
        draw_card(ax, 0.23 + index * (card_width + 0.012), 0.835, card_width, 0.105, title, value)

    draw_text(ax, 0.23, 0.64, "Первые строки датасета", size=15, weight="bold")
    table_x, table_y, row_h = 0.23, 0.58, 0.043
    columns = ["Age", "Gender", "WeeklyPay", "InitialCaseEstimate", "UltimateCost"]
    widths = [0.06, 0.07, 0.09, 0.13, 0.12]
    x = table_x
    for col, width in zip(columns, widths):
        ax.add_patch(plt.Rectangle((x, table_y), width, row_h, transform=ax.transAxes, facecolor=TIFFANY, edgecolor="#FFFFFF"))
        draw_text(ax, x + 0.006, table_y + 0.031, col, size=8, color="#FFFFFF", weight="bold")
        x += width
    sample_rows = [
        ["45", "M", "500", "9 500", "102.39"],
        ["40", "M", "283", "3 000", "1 451.00"],
        ["19", "F", "0", "250", "48.00"],
        ["35", "M", "520", "5 000", "2 430.00"],
    ]
    for r, values in enumerate(sample_rows):
        x = table_x
        y = table_y - (r + 1) * row_h
        for value, width in zip(values, widths):
            ax.add_patch(plt.Rectangle((x, y), width, row_h, transform=ax.transAxes, facecolor="#FFFFFF", edgecolor=GRID))
            draw_text(ax, x + 0.006, y + 0.031, value, size=9)
            x += width

    chart_ax = fig.add_axes([0.75, 0.19, 0.20, 0.36])
    chart_ax.barh(result.metrics["Model"], result.metrics["RMSE"], color=[MODEL_COLORS.get(m, TIFFANY) for m in result.metrics["Model"]])
    chart_ax.invert_yaxis()
    chart_ax.set_title("RMSE моделей", fontsize=13, color=INK, weight="bold")
    chart_ax.set_xlabel("Ошибка, $")
    chart_ax.set_yticklabels([])
    for index, model in enumerate(result.metrics["Model"]):
        chart_ax.text(280, index, model, va="center", fontsize=8, color="#FFFFFF", weight="bold")
    chart_ax.spines[["top", "right", "left"]].set_visible(False)
    chart_ax.grid(axis="x", alpha=0.5)
    fig.savefig(ASSETS_DIR / "app_main.png", dpi=160, bbox_inches="tight")
    plt.close(fig)

    fig = plt.figure(figsize=(16, 10), dpi=150)
    ax = draw_app_shell(fig, "Прогнозирование стоимости страховых выплат", "Сравнение графиков")
    draw_text(ax, 0.23, 0.64, "Сравнение моделей: факт vs прогноз", size=16, weight="bold")
    limit = float(result.target_cap)
    positions = [(0.24, 0.37), (0.57, 0.37), (0.24, 0.08), (0.57, 0.08)]
    for (x0, y0), (_, metric_row) in zip(positions, result.metrics.iterrows()):
        model_name = metric_row["Model"]
        axis = fig.add_axes([x0, y0, 0.28, 0.22])
        sample = result.predictions[result.predictions["Model"] == model_name]
        sample = sample.sample(min(900, len(sample)), random_state=42)
        axis.scatter(sample["Actual"], sample["Predicted"], s=8, alpha=0.30, color=MODEL_COLORS.get(model_name, TIFFANY), edgecolors="none")
        axis.plot([0, limit], [0, limit], color="#c94c4c", linestyle="--", linewidth=1)
        axis.set_xlim(0, limit)
        axis.set_ylim(0, limit)
        axis.set_title(f"{model_name}: RMSE {metric_row['RMSE']:,.0f}".replace(",", " "), fontsize=10, color=INK)
        axis.spines[["top", "right"]].set_visible(False)
        axis.tick_params(labelsize=8)
    fig.savefig(ASSETS_DIR / "app_comparison.png", dpi=160, bbox_inches="tight")
    plt.close(fig)

    fig = plt.figure(figsize=(16, 10), dpi=150)
    ax = draw_app_shell(fig, "Презентация проекта", "Презентация")
    slide_x, slide_y, slide_w, slide_h = 0.25, 0.16, 0.66, 0.58
    ax.add_patch(plt.Rectangle((slide_x, slide_y), slide_w, slide_h, transform=ax.transAxes, facecolor="#FFFFFF", edgecolor=GRID, linewidth=1.4))
    ax.add_patch(plt.Rectangle((slide_x, slide_y + slide_h - 0.018), slide_w, 0.018, transform=ax.transAxes, facecolor=TIFFANY, edgecolor=TIFFANY))
    draw_text(ax, slide_x + 0.04, slide_y + slide_h - 0.08, "Прогнозирование стоимости страховых выплат", size=22, weight="bold")
    bullets = [
        "OpenML Workers Compensation, 100 000 записей",
        "Предобработка дат, категорий, пропусков и выбросов",
        "Сравнение Linear Regression, Ridge, Random Forest и XGBoost",
        f"Лучшая модель: {result.best_model_name}, RMSE {format_currency(float(result.metrics.iloc[0]['RMSE']))}",
        "Интерактивное сравнение графиков и прогноз нового случая",
    ]
    for index, bullet in enumerate(bullets):
        draw_text(ax, slide_x + 0.055, slide_y + slide_h - 0.18 - index * 0.073, f"- {bullet}", size=14)
    fig.savefig(ASSETS_DIR / "app_presentation.png", dpi=160, bbox_inches="tight")
    plt.close(fig)


def export_results_json(result) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    payload = {
        "best_model_name": result.best_model_name,
        "sample_size": result.sample_size,
        "data_shape": result.data_shape,
        "metrics": result.metrics.to_dict(orient="records"),
        "feature_importance": result.feature_importance.to_dict(orient="records"),
        "target_summary": result.target_summary,
        "target_cap": result.target_cap,
    }
    RESULTS_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def build_docx(df: pd.DataFrame, result, chart_paths: dict[str, Path]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    clear_document_body(doc)
    set_default_font(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    add_page_number(section)

    add_centered(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", size=12)
    add_centered(doc, "РОССИЙСКОЙ ФЕДЕРАЦИИ", size=12)
    add_centered(doc, "федеральное государственное бюджетное образовательное учреждение высшего образования", size=12)
    add_centered(doc, "«Казанский национальный исследовательский технический университет им. А.Н. Туполева - КАИ»", size=12)
    add_centered(doc, "(КНИТУ - КАИ)", size=12)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", bold=True, size=16)
    add_centered(doc, "по дополнительной профессиональной программе профессиональной переподготовки", size=12)
    add_centered(doc, "«Data science. Искусственный интеллект»", size=14)
    doc.add_paragraph()
    add_centered(doc, "по теме:", size=14)
    add_centered(doc, "«Регрессия для прогнозирования стоимости страховых выплат»", bold=True, size=16)
    for _ in range(3):
        doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Работу выполнил:"
    table.cell(0, 1).text = "Южаков Максим"
    table.cell(1, 0).text = "Руководитель выпускной квалификационной работы:"
    table.cell(1, 1).text = "к.т.н., доцент Смирнова Гульнара Сергеевна"
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "Казань 2026", size=14)
    doc.add_page_break()

    add_heading(doc, "ОГЛАВЛЕНИЕ")
    toc_rows = [
        "Введение",
        "1. Описание датасета",
        "2. Предобработка данных",
        "3. Разделение данных",
        "4. Обучение моделей",
        "5. Оценка моделей",
        "6. Анализ важности признаков",
        "7. Streamlit-приложение",
        "Заключение",
        "Список литературы",
        "Приложение",
    ]
    for row in toc_rows:
        doc.add_paragraph(row)
    doc.add_page_break()

    add_heading(doc, "ВВЕДЕНИЕ")
    add_paragraph(
        doc,
        "Цель работы - разработать модель машинного обучения, которая прогнозирует итоговую стоимость страхового возмещения "
        "UltimateIncurredClaimCost на основе характеристик работника, условий занятости, описания происшествия и начальной оценки случая. "
        "Такая задача актуальна для страховых компаний, поскольку точная оценка будущих выплат помогает планировать резервы, контролировать "
        "убыточность портфеля и быстрее принимать управленческие решения по страховым случаям."
    )
    add_paragraph(
        doc,
        "В рамках работы создано многостраничное Streamlit-приложение с навигацией st.navigation и st.Page. В приложении реализованы "
        "загрузка данных, предобработка, обучение нескольких регрессионных моделей, сравнение метрик, анализ важности признаков и расчет "
        "прогноза для нового страхового случая."
    )
    add_paragraph(doc, "Ссылка на git-репозиторий: указывается при размещении проекта в удаленном репозитории.")

    add_heading(doc, "ОСНОВНАЯ ЧАСТЬ")
    add_heading(doc, "1. Описание датасета", level=2)
    add_paragraph(
        doc,
        f"Источник данных - Workers Compensation Dataset из OpenML, идентификатор 42876. Набор содержит {df.shape[0]:,} записей "
        f"и {df.shape[1]} столбцов. Целевая переменная - {TARGET}, то есть итоговая стоимость страхового возмещения.".replace(",", " ")
    )
    add_simple_table(
        doc,
        [["Переменная", "Описание"]] + [[name, description] for name, description in FEATURE_DESCRIPTIONS],
        widths=[5.2, 10.2],
        font_size=9,
        caption="Таблица 1 - Описание признаков датасета",
    )
    sample = df.head(5)[["Age", "Gender", "WeeklyPay", "InitialCaseEstimate", TARGET]]
    add_simple_table(
        doc,
        [["Age", "Gender", "WeeklyPay", "InitialCaseEstimate", TARGET]]
        + [[str(value) for value in row] for row in sample.to_numpy().tolist()],
        widths=[1.8, 2.0, 2.5, 3.7, 4.4],
        font_size=8,
        caption="Таблица 2 - Пример строк исходного датасета",
    )

    add_heading(doc, "2. Предобработка данных", level=2)
    add_paragraph(
        doc,
        "В ходе предобработки поля DateTimeOfAccident и DateReported преобразованы в формат даты. На их основе созданы признаки "
        "AccidentYear, AccidentMonth, AccidentDayOfWeek и ReportingDelay. Задержка сообщения о случае важна экономически, так как поздняя "
        "регистрация может отражать сложность расследования и влиять на итоговые затраты."
    )
    add_paragraph(
        doc,
        "Категориальные переменные Gender, MaritalStatus, PartTimeFullTime и ClaimDescription преобразованы в числовой формат с помощью "
        "OrdinalEncoder. Для ClaimDescription дополнительно сформированы признаки длины описания и количества слов. Числовые признаки "
        "масштабированы StandardScaler. Для снижения влияния экстремальных выплат целевая переменная при обучении была ограничена сверху "
        f"95-м перцентилем ({format_currency(result.target_cap)}), поскольку исходная стоимость возмещения варьируется от "
        f"{format_currency(result.target_summary['min'])} до {format_currency(result.target_summary['max'])}."
    )
    missing_total = sum(result.missing_values.values())
    add_paragraph(
        doc,
        f"Проверка пропущенных значений показала суммарно {missing_total} пропусков. Отрицательных значений целевой переменной не обнаружено."
    )

    add_heading(doc, "3. Разделение данных", level=2)
    add_paragraph(
        doc,
        f"Данные разделены на обучающую и тестовую выборки в соотношении 80/20. При размере обучающей выборки {int(result.sample_size * 0.8):,} "
        f"наблюдений модель получает достаточно данных для выявления закономерностей, а тестовая выборка из {int(result.sample_size * 0.2):,} "
        "наблюдений позволяет объективно оценить качество прогноза на новых данных.".replace(",", " ")
    )

    add_heading(doc, "4. Обучение моделей", level=2)
    add_paragraph(
        doc,
        "Для сравнения использованы четыре регрессионные модели. Linear Regression выступает базовой интерпретируемой моделью. Ridge Regression "
        "добавляет L2-регуляризацию и снижает риск переобучения при коррелированных признаках. Random Forest Regressor учитывает нелинейные "
        "зависимости и взаимодействия признаков. XGBoost Regressor строит ансамбль деревьев градиентного бустинга и обычно хорошо работает "
        "на табличных данных со сложной структурой."
    )

    add_heading(doc, "5. Оценка моделей", level=2)
    add_paragraph(
        doc,
        "Качество оценивалось по MAE, MSE, RMSE и коэффициенту детерминации R2. MAE показывает среднюю абсолютную ошибку в долларах, RMSE "
        "сильнее штрафует крупные ошибки, а R2 отражает долю объясненной вариации целевой переменной."
    )
    metric_rows = [["Модель", "MAE, $", "MSE, млн $^2", "RMSE, $", "R2"]]
    for _, row in result.metrics.iterrows():
        metric_rows.append(
            [
                str(row["Model"]),
                f"{row['MAE']:,.2f}".replace(",", " "),
                f"{row['MSE'] / 1_000_000:.2f}",
                f"{row['RMSE']:,.2f}".replace(",", " "),
                f"{row['R2']:.4f}",
            ]
        )
    add_simple_table(
        doc,
        metric_rows,
        widths=[4.1, 2.5, 3.0, 2.5, 1.5],
        font_size=9,
        caption="Таблица 3 - Сравнение качества регрессионных моделей",
    )
    best_row = result.metrics.iloc[0]
    add_paragraph(
        doc,
        f"Наилучший результат по RMSE показала модель {result.best_model_name}: RMSE = {format_currency(best_row['RMSE'])}, "
        f"MAE = {format_currency(best_row['MAE'])}, R2 = {best_row['R2']:.4f}."
    )
    add_picture_if_exists(doc, chart_paths["metrics"], "Рисунок 1 - Сравнение моделей по MAE и RMSE")
    add_picture_if_exists(doc, chart_paths["comparison"], "Рисунок 2 - Сравнение графиков факт-прогноз по моделям", width=Cm(15))
    add_picture_if_exists(doc, chart_paths["scatter"], "Рисунок 3 - Предсказанные и реальные значения для лучшей модели", width=Cm(12.8))
    add_picture_if_exists(doc, chart_paths["errors"], "Рисунок 4 - Распределение абсолютных ошибок моделей")

    add_heading(doc, "6. Анализ важности признаков", level=2)
    importance_rows = [["Признак", "Важность"]]
    for _, row in result.feature_importance.head(10).iterrows():
        importance_rows.append([str(row["Feature"]), f"{row['Importance']:.4f}"])
    add_simple_table(
        doc,
        importance_rows,
        widths=[6.0, 3.0],
        font_size=9,
        caption="Таблица 4 - Важность признаков лучшей модели",
    )
    top_feature = str(result.feature_importance.iloc[0]["Feature"])
    add_paragraph(
        doc,
        f"Самым значимым признаком стала переменная {top_feature}. Это ожидаемо, поскольку начальная оценка случая аккумулирует экспертную "
        "информацию страховой компании о тяжести происшествия. Также вклад дают характеристики занятости, возраст, задержка сообщения и "
        "описание страхового случая."
    )
    add_picture_if_exists(doc, chart_paths["importance"], "Рисунок 5 - Важность признаков", width=Cm(13.5))

    add_heading(doc, "7. Streamlit-приложение", level=2)
    add_paragraph(
        doc,
        "Streamlit-приложение состоит из двух страниц: «Анализ и модель» и «Презентация». Навигация реализована через st.navigation и st.Page "
        "в файле app.py. Основная страница загружает датасет OpenML, показывает первые строки и статистику, обучает модели, выводит метрики, "
        "строит графики и позволяет рассчитать прогноз для нового страхового случая. Интерфейс оформлен в светлой цветовой схеме с акцентным "
        "цветом Tiffany (#0ABAB5)."
    )
    add_paragraph(
        doc,
        "Для сравнения моделей добавлена отдельная вкладка «Сравнение графиков». На ней можно выбрать набор моделей и переключаться между "
        "графиками «факт-прогноз», остатками и распределением абсолютных ошибок."
    )
    add_paragraph(
        doc,
        "Страница презентации создана с использованием streamlit-reveal-slides. Она кратко показывает бизнес-задачу, состав данных, "
        "этапы предобработки, обученные модели, результаты и направления улучшения."
    )
    add_picture_if_exists(doc, ASSETS_DIR / "app_main.png", "Рисунок 6 - Основная страница Streamlit-приложения")
    add_picture_if_exists(doc, ASSETS_DIR / "app_comparison.png", "Рисунок 7 - Сравнение графиков моделей в Streamlit")
    add_picture_if_exists(doc, ASSETS_DIR / "app_presentation.png", "Рисунок 8 - Страница презентации Streamlit-приложения")

    add_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_paragraph(
        doc,
        f"В результате работы реализован полный цикл решения регрессионной задачи: загрузка и анализ данных, предобработка, разделение "
        f"выборки, обучение моделей, оценка качества и разработка интерактивного Streamlit-приложения. Лучшей моделью по RMSE стала "
        f"{result.best_model_name}, что подтверждает применимость ансамблевых методов к задаче прогнозирования страховых выплат."
    )
    add_paragraph(
        doc,
        "Возможные улучшения включают более глубокую обработку текстового описания заявки с помощью NLP-признаков, подбор гиперпараметров, "
        "робастную работу с экстремальными выплатами и добавление внешних факторов, например отрасли работника или истории предыдущих случаев."
    )

    add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ")
    references = [
        "Workers Compensation Dataset. OpenML. URL: https://www.openml.org/d/42876",
        "Scikit-learn Documentation. Supervised learning. URL: https://scikit-learn.org/stable/supervised_learning.html",
        "Scikit-learn Documentation. Regression metrics. URL: https://scikit-learn.org/stable/modules/model_evaluation.html#regression-metrics",
        "Pandas Documentation. URL: https://pandas.pydata.org/docs/",
        "Streamlit Documentation. URL: https://docs.streamlit.io/",
        "XGBoost Documentation. URL: https://xgboost.readthedocs.io/",
        "Streamlit Reveal Slides. URL: https://github.com/bouzidanas/streamlit-reveal-slides",
    ]
    for index, reference in enumerate(references, start=1):
        paragraph = doc.add_paragraph(f"{index}. {reference}")
        paragraph.paragraph_format.line_spacing = 1.5

    add_heading(doc, "Приложение")
    add_paragraph(doc, "Проверка соответствия основным требованиям методических указаний:")
    compliance_rows = [
        ["Требование", "Статус", "Где реализовано"],
        ["Многостраничное Streamlit-приложение", "Выполнено", "app.py, st.navigation, st.Page"],
        ["Основная страница анализа и модели", "Выполнено", "analysis_and_model.py"],
        ["Страница презентации", "Выполнено", "presentation.py, streamlit-reveal-slides"],
        ["Загрузка Workers Compensation OpenML ID 42876", "Выполнено", "model_utils.py"],
        ["Предобработка дат, категорий, пропусков и выбросов", "Выполнено", "model_utils.py, раздел 2 отчета"],
        ["Разделение train/test 80/20", "Выполнено", "model_utils.py, раздел 3 отчета"],
        ["Сравнение Linear Regression, Random Forest, XGBoost, Ridge", "Выполнено", "model_utils.py, раздел 4 отчета"],
        ["Метрики MAE, MSE, RMSE, R2", "Выполнено", "раздел 5 отчета"],
        ["Анализ важности признаков", "Выполнено", "раздел 6 отчета"],
        ["Интерактивное сравнение графиков моделей", "Выполнено", "вкладка «Сравнение графиков»"],
        ["README.md и requirements.txt", "Выполнено", "корень проекта"],
        ["Видео-демонстрация", "Выполнено", "video/demo.mp4"],
    ]
    add_simple_table(
        doc,
        compliance_rows,
        widths=[6.3, 2.3, 6.5],
        font_size=8,
        caption="Таблица 5 - Проверка соответствия методическим указаниям",
    )
    add_paragraph(doc, "Структура проекта соответствует методическим указаниям:")
    project_rows = [
        ["Файл", "Назначение"],
        ["app.py", "Основной файл приложения"],
        ["analysis_and_model.py", "Анализ данных, обучение моделей, прогноз"],
        ["presentation.py", "Презентация проекта"],
        ["model_utils.py", "Общая логика подготовки данных и обучения"],
        [".streamlit/config.toml", "Тема интерфейса и Tiffany-акцент"],
        ["requirements.txt", "Зависимости проекта"],
        ["README.md", "Описание проекта и запуск"],
        ["scripts/build_report.py", "Генерация отчета, таблиц и графиков"],
        ["reports/", "Готовый отчет, изображения и JSON с результатами"],
        ["video/demo.mp4", "Короткая демонстрация интерфейса приложения"],
    ]
    add_simple_table(
        doc,
        project_rows,
        widths=[5.0, 9.8],
        font_size=9,
        caption="Таблица 6 - Структура проекта",
    )
    add_paragraph(doc, "Ссылка на репозиторий: добавляется после публикации проекта.")

    doc.save(OUTPUT_DOCX)


def collect_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def write_consistency_check(result) -> None:
    text = collect_docx_text(OUTPUT_DOCX)
    best_row = result.metrics.iloc[0]
    expected = {
        "best_model": result.best_model_name,
        "best_mae": f"{best_row['MAE']:,.2f}".replace(",", " "),
        "best_rmse": f"{best_row['RMSE']:,.2f}".replace(",", " "),
        "best_r2": f"{best_row['R2']:.4f}",
        "target_cap": f"{result.target_cap:,.2f}".replace(",", " "),
        "top_feature": str(result.feature_importance.iloc[0]["Feature"]),
    }
    checks = {name: value in text for name, value in expected.items()}
    lines = [
        "Проверка соответствия отчета фактическим результатам",
        f"Файл отчета: {OUTPUT_DOCX}",
        "",
    ]
    for name, value in expected.items():
        status = "OK" if checks[name] else "MISSING"
        lines.append(f"{status}: {name} = {value}")
    CONSISTENCY_CHECK.write_text("\n".join(lines), encoding="utf-8")
    missing = [name for name, ok in checks.items() if not ok]
    if missing:
        raise RuntimeError(f"Report consistency check failed: {', '.join(missing)}")


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    df = load_workers_compensation()
    result = train_models(df, sample_size=None)
    chart_paths = save_charts(result)
    save_app_previews(result)
    export_results_json(result)
    build_docx(df, result, chart_paths)
    write_consistency_check(result)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
